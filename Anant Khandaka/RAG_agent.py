import asyncio
import os
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
import hashlib

import aiofiles
import aiohttp
from autogen_agentchat.agents import AssistantAgent
from autogen_agentchat.ui import Console
from autogen_core.memory import Memory, MemoryContent, MemoryMimeType
from autogen_ext.memory.chromadb import ChromaDBVectorMemory, PersistentChromaDBVectorMemoryConfig
from autogen_ext.models.ollama import OllamaChatCompletionClient


try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


@dataclass
class Document:
    """Represents a document with metadata."""
    content: str
    source: str
    doc_type: str
    hash: str
    page_info: Optional[str] = None


class AdvancedDocumentParser:
    """Handles parsing of multiple document types."""

    @staticmethod
    async def parse_document(source: str) -> Document:
        """Parse document from various sources and formats."""
        content = ""
        doc_type = "unknown"
        
        if source.startswith(("http://", "https://")):
            content = await AdvancedDocumentParser._fetch_url(source)
            doc_type = "url"
        elif source.lower().endswith(".pdf") and PDF_SUPPORT:
            content, doc_type = AdvancedDocumentParser._parse_pdf(source)
        elif source.lower().endswith(".docx") and DOCX_SUPPORT:
            content, doc_type = AdvancedDocumentParser._parse_docx(source)
        elif source.lower().endswith(".txt"):
            content = await AdvancedDocumentParser._fetch_file(source)
            doc_type = "txt"
        else:
            content = await AdvancedDocumentParser._fetch_file(source)
            doc_type = "text"
        
        doc_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        return Document(content=content, source=source, doc_type=doc_type, hash=doc_hash)

    @staticmethod
    async def _fetch_url(url: str) -> str:
        """Fetch content from URL."""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=10) as response:
                    text = await response.text()
                    return AdvancedDocumentParser._clean_html(text)
        except Exception as e:
            print(f"Error fetching {url}: {e}")
            return ""

    @staticmethod
    async def _fetch_file(path: str) -> str:
        """Fetch content from file."""
        try:
            async with aiofiles.open(path, "r", encoding="utf-8") as f:
                return await f.read()
        except Exception as e:
            print(f"Error reading {path}: {e}")
            return ""

    @staticmethod
    def _parse_pdf(path: str) -> Tuple[str, str]:
        """Extract text from PDF with page info."""
        try:
            text = ""
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text += f"\n[Page {page_num}]\n{page_text}\n"
            return text, "pdf"
        except Exception as e:
            print(f"Error parsing PDF {path}: {e}")
            return "", "pdf"

    @staticmethod
    def _parse_docx(path: str) -> Tuple[str, str]:
        """Extract text from DOCX."""
        try:
            doc = Document(path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return text, "docx"
        except Exception as e:
            print(f"Error parsing DOCX {path}: {e}")
            return "", "docx"

    @staticmethod
    def _clean_html(html: str) -> str:
        """Remove HTML tags and clean whitespace."""
        text = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL)
        text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()


class AdvancedChunker:
    """Advanced text chunking strategies for optimal RAG performance."""

    def __init__(self, chunk_size: int = 600, overlap: int = 150, smart_split: bool = True):
        """
        Args:
            chunk_size: Target chunk size in characters
            overlap: Overlap between chunks for context preservation
            smart_split: Use intelligent splitting at sentence boundaries
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.smart_split = smart_split

    def chunk_text(self, text: str, doc_source: str = "") -> List[Dict]:
        """
        Advanced chunking with multiple strategies.
        Returns chunks with metadata for better retrieval.
        """
        text = self._normalize_text(text)
        
        if len(text) < self.chunk_size:
            return [{
                "content": text,
                "source": doc_source,
                "chunk_index": 0,
                "chunk_size": len(text),
            }]
        
        chunks = self._semantic_chunking(text)
        
        return [{
            "content": chunk,
            "source": doc_source,
            "chunk_index": idx,
            "chunk_size": len(chunk),
        } for idx, chunk in enumerate(chunks)]

    def _normalize_text(self, text: str) -> str:
        """Normalize text for better chunking."""
        
        text = re.sub(r"\n\n+", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        lines = text.split("\n")
        lines = [l for l in lines if len(l.strip()) > 5]
        return "\n".join(lines).strip()

    def _semantic_chunking(self, text: str) -> List[str]:
        """Split text at sentence boundaries for better semantic chunks."""
        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            test_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        if self.overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks)
        
        return chunks

    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences intelligently."""
        text = re.sub(r"([.!?])\s+", r"\1|", text)
        sentences = text.split("|")
        return [s.strip() for s in sentences if s.strip()]

    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """Add overlap between chunks for context."""
        if len(chunks) <= 1:
            return chunks
        
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1]
            curr_chunk = chunks[i]
            
            overlap_text = " ".join(prev_chunk.split()[-5:])
            combined = overlap_text + " " + curr_chunk
            overlapped.append(combined)
        
        return overlapped


class OptimizedRAGIndexer:
    """Optimized document indexer with advanced chunking and caching."""

    def __init__(self, memory: Memory, chunk_size: int = 600, cache_dir: Optional[str] = None):
        self.memory = memory
        self.chunker = AdvancedChunker(chunk_size=chunk_size, overlap=150)
        self.cache_dir = cache_dir or os.path.join(str(Path.home()), ".rag_cache")
        Path(self.cache_dir).mkdir(exist_ok=True)
        self.indexed_docs = set()

    def _get_cache_path(self, doc_hash: str) -> Path:
        """Get cache file path for document."""
        return Path(self.cache_dir) / f"{doc_hash}.cache"

    async def index_documents(self, sources: List[str]) -> Tuple[int, int]:
        """
        Index documents with smart caching.
        Returns (total_chunks, total_documents)
        """
        total_chunks = 0
        doc_count = 0

        for source in sources:
            try:
                doc = await AdvancedDocumentParser.parse_document(source)
                
                if not doc.content.strip():
                    print(f"Empty document: {source}")
                    continue
                
                cache_path = self._get_cache_path(doc.hash)
                if cache_path.exists() and doc.hash in self.indexed_docs:
                    print(f"Using cached: {source}")
                    chunk_count = int(cache_path.read_text().strip().split("\n")[0])
                    total_chunks += chunk_count
                    doc_count += 1
                    continue
                
                chunks = self.chunker.chunk_text(doc.content, doc_source=source)
                print(f"  └─ {len(chunks)} chunks from {source}")
                
                for chunk_data in chunks:
                    await self.memory.add(
                        MemoryContent(
                            content=chunk_data["content"],
                            mime_type=MemoryMimeType.TEXT,
                            metadata={
                                "source": doc.source,
                                "doc_type": doc.doc_type,
                                "chunk_index": chunk_data["chunk_index"],
                                "doc_hash": doc.hash,
                            }
                        )
                    )
                
                cache_path.write_text(f"{len(chunks)}\n")
                self.indexed_docs.add(doc.hash)
                
                total_chunks += len(chunks)
                doc_count += 1

            except Exception as e:
                print(f"Error indexing {source}: {e}")

        return total_chunks, doc_count


class RAGAgentFactory:
    """Factory for creating optimized RAG agents."""

    @staticmethod
    def create_rag_system_prompt() -> str:
        """Create optimized system prompt for RAG."""
        return """You are an expert document analyzer. Your task is to answer user questions accurately and comprehensively based on the provided document content.

CRITICAL RULES - FOLLOW STRICTLY:
1. ONLY output natural language answers - NO code, NO objects, NO data structures
2. NEVER output MemoryContent, metadata, or technical details
3. NEVER show Python objects, lists, or internal system information
4. ALWAYS answer based on document content - do not apologize or say you cannot answer
5. Provide direct, factual answers without unnecessary preamble
6. If asked "what is this document about?" - provide a clear summary in 2-3 sentences
7. If asked for lists (skills, projects, etc.) - return formatted lists with clear categorization
8. Extract specific information accurately from the provided context
9. Be concise but thorough
10. Format lists with proper structure and categories
11. For resume/CV documents, identify and extract: skills, experience, education, projects, certifications
12. Return ONLY clean, readable information that's easy to read

Examples of WRONG output:
- [MemoryContent(content='...' mime_type='...')]
- MemoryMimeType.TEXT
- metadata={'doc_hash': '...'}
- Any Python objects or internal structures

Examples of CORRECT output:
- "The person's skills include Python, FastAPI, and Machine Learning"
- "Educational background: Bachelor of Technology in Computer Science"
- Natural, conversational English sentences

REMEMBER: Your output will be displayed to users. Make it clean, professional, and human-readable."""

    @staticmethod
    async def create_agent(
        model: str = "llama3.2:1b",
        memory: Optional[Memory] = None,
        collection_name: str = "rag_documents"
    ) -> Tuple[AssistantAgent, ChromaDBVectorMemory]:
        """Create optimized RAG agent with Ollama."""
        
        if memory is None:
            memory = ChromaDBVectorMemory(
                config=PersistentChromaDBVectorMemoryConfig(
                    collection_name=collection_name,
                    persistence_path=os.path.join(str(Path.home()), f".chromadb_{collection_name}"),
                    k=20,  
                    score_threshold=0.15,  
                )
            )
        
        agent = AssistantAgent(
            name="rag_analyst",
            model_client=OllamaChatCompletionClient(
                model=model,
                temperature=0.2,  
                top_p=0.7,
            ),
            memory=[memory],
            system_message=RAGAgentFactory.create_rag_system_prompt()
        )
        
        return agent, memory


async def main() -> None:
    """Main RAG application."""
    
    print("\n" + "="*70)
    print("ADVANCED RAG AGENT WITH OLLAMA - OPTIMAL PERFORMANCE")
    print("="*70 + "\n")
    
    DOCUMENT_SOURCES = [
        "Agents/Anant_2.pdf",
    ]
    
    TEST_QUESTIONS = [
        "Who is this person and what is their current or most recent job?",
        "What are this person's main technical skills and programming languages?",
        "What projects has this person worked on and what technologies were used?",
        "What improvements or achievements did this person achieve in their internships?",
        "What educational background and certifications does this person have?",
    ]
    
    OLLAMA_MODEL = "llama3.2:1b"  
    print("System Configuration:")
    print(f"  Model: {OLLAMA_MODEL}")
    print(f"  Documents: {len(DOCUMENT_SOURCES)}")
    print(f"  Questions: {len(TEST_QUESTIONS)}\n")
    
    available_docs = []
    for src in DOCUMENT_SOURCES:
        if src.startswith(("http://", "https://")) or Path(src).exists():
            available_docs.append(src)
    
    if not available_docs:
        print("No documents found. Using demo mode.")
        available_docs = [
            "https://raw.githubusercontent.com/microsoft/autogen/main/README.md"
        ]
    
    print("Initializing RAG system...")
    rag_agent, rag_memory = await RAGAgentFactory.create_agent(
        model=OLLAMA_MODEL,
        collection_name="ehr_documents"
    )
    
    await rag_memory.clear()
    
    print(f"\nIndexing {len(available_docs)} document(s)...")
    indexer = OptimizedRAGIndexer(memory=rag_memory, chunk_size=800) 
    total_chunks, doc_count = await indexer.index_documents(available_docs)
    
    print(f"\nIndexing complete!")
    print(f"  Documents processed: {doc_count}")
    print(f"  Total chunks: {total_chunks}")
    print(f"  Retrieval: top-20 chunks with similarity threshold 0.15")
    print(f"  Model: {OLLAMA_MODEL} (better for RAG than qwen2.5:3b)\n")
    
    
    for i, question in enumerate(TEST_QUESTIONS, 1):
        print(f"\nQuestion {i}/{len(TEST_QUESTIONS)}: {question}")
        print("-"*70)
        
        try:
            stream = rag_agent.run_stream(task=question)
            await Console(stream)
        except Exception as e:
            print(f"Error: {e}")
        
        print()
    
    # Cleanup
    await rag_memory.close()


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n\nRAG Agent interrupted by user")
    except Exception as e:
        print(f"\n\nFatal error: {e}")
